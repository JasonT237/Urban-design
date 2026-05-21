from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUTPUT = "docs/Urban_Sanctuary_Beginner_Guide.docx"
PRIMARY = RGBColor(8, 76, 112)
DARK = RGBColor(15, 23, 42)
MUTED = RGBColor(71, 85, 105)
SOFT = "F7F8F0"
HEADER = "E6E8DC"


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text, bold=False, color=None):
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(str(text))
    run.bold = bold
    run.font.size = Pt(9)
    if color:
        run.font.color.rgb = color
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_table(doc, headers, rows):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for idx, header in enumerate(headers):
        shade_cell(table.rows[0].cells[idx], HEADER)
        set_cell_text(table.rows[0].cells[idx], header, bold=True, color=PRIMARY)
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            set_cell_text(cells[idx], value)
    doc.add_paragraph()
    return table


def add_callout(doc, title, body):
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    cell = table.cell(0, 0)
    shade_cell(cell, SOFT)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(title)
    run.bold = True
    run.font.color.rgb = PRIMARY
    run.font.size = Pt(10)
    p2 = cell.add_paragraph()
    p2.paragraph_format.space_after = Pt(0)
    r2 = p2.add_run(body)
    r2.font.size = Pt(9)
    r2.font.color.rgb = MUTED
    doc.add_paragraph()


def heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        run.font.color.rgb = PRIMARY if level <= 2 else DARK
    return p


def para(doc, text, bold=False, italic=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.08
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.color.rgb = MUTED
    run.font.size = Pt(10.5)
    return p


def bullet(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(f"- {text}")
    run.font.size = Pt(10)
    run.font.color.rgb = MUTED


def number(doc, index, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(f"{index}. {text}")
    run.font.size = Pt(10)
    run.font.color.rgb = MUTED


def code_block(doc, text):
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    cell = table.cell(0, 0)
    shade_cell(cell, "F3F4F6")
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(text)
    run.font.name = "Consolas"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Consolas")
    run.font.size = Pt(8.5)
    doc.add_paragraph()


def configure_document(doc):
    section = doc.sections[0]
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.7)
    section.right_margin = Inches(0.7)

    styles = doc.styles
    styles["Normal"].font.name = "Aptos"
    styles["Normal"].font.size = Pt(10.5)
    styles["Normal"].font.color.rgb = MUTED
    for name, size in [("Heading 1", 20), ("Heading 2", 15), ("Heading 3", 12)]:
        styles[name].font.name = "Aptos Display"
        styles[name].font.size = Pt(size)
        styles[name].font.bold = True
        styles[name].font.color.rgb = PRIMARY


def cover(doc):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run("Urban Sanctuary Booking Portal")
    r.bold = True
    r.font.size = Pt(28)
    r.font.color.rgb = PRIMARY

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Beginner Developer Guide")
    r.font.size = Pt(16)
    r.font.color.rgb = DARK

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("How the app works, what every folder does, and where to make future changes.")
    r.font.size = Pt(11)
    r.font.color.rgb = MUTED

    doc.add_paragraph()
    add_callout(
        doc,
        "How to use this guide",
        "Read sections 1 to 5 for the big picture. Use the route map, API map, hook map, and change recipes when you need to modify the app yourself.",
    )
    add_callout(
        doc,
        "Current implementation summary",
        "The public property list, booking flow, payment initiation, profile, booking history, cancel booking, and several admin pages now talk to the backend. Saved places, some marketing content, and some admin content-editor controls are still static or partly placeholder.",
    )
    doc.add_page_break()


def build_doc():
    doc = Document()
    configure_document(doc)
    cover(doc)

    heading(doc, "1. What The App Does")
    para(
        doc,
        "Urban Sanctuary is a React single-page booking portal for short-stay apartments in Douala. Guests browse properties, create bookings, pay, view history, cancel eligible bookings, and manage their profile. Admins use a separate /admin portal to manage users, listings, bookings, support tickets, permissions, and system settings.",
    )
    heading(doc, "Main User Roles", 2)
    add_table(
        doc,
        ["Role", "What they can do", "Where they go after login"],
        [
            ["Visitor", "Browse Discover, public listings, apartment details, and support.", "No login needed"],
            ["Guest", "Book apartments, pay, view dashboard/history/profile, cancel eligible bookings.", "/discover"],
            ["Admin", "Use the admin console and see system-wide resources.", "/admin"],
        ],
    )

    heading(doc, "2. Tech Stack")
    add_table(
        doc,
        ["Tool", "Purpose"],
        [
            ["React 19", "Builds the component-based user interface."],
            ["React Router 7", "Maps URLs like /apartments and /admin/users to components."],
            ["Vite 8", "Runs the local dev server and production build."],
            ["Tailwind CSS 3", "Styles the UI with utility classes inside JSX."],
            ["Fetch", "The real request engine used by src/services/apiClient.js."],
            ["Axios", "Installed in package.json, but the current shared API layer uses fetch."],
        ],
    )
    heading(doc, "Environment Setup", 2)
    code_block(doc, "VITE_API_BASE_URL=https://sanaps.mooo.com/api/v1")
    for index, step in enumerate([
        "Run npm install from C:\\Users\\Utilisateur\\booking-portal.",
        "Run npm run dev to start Vite.",
        "Open the printed localhost URL, usually http://localhost:5173.",
        "Run npm run lint before pushing code.",
        "Run npm run build before deployment.",
    ], start=1):
        number(doc, index, step)

    doc.add_page_break()
    heading(doc, "3. Folder Map")
    code_block(
        doc,
        """booking-portal/
  docs/        Word guides and swagger.yaml
  scripts/     Documentation generator scripts
  src/
    routes/    AppRouter and route protection
    pages/     Public and guest route screens
    components/ Reusable UI and feature component folders
    hooks/     Custom hooks for state, API loading, tokens
    services/  API request wrappers
    context/   Shared providers
    lib/       Adapters, formatters, filters, constants
    data/      Remaining static sample content
""",
    )
    add_callout(
        doc,
        "Simple rule",
        "Pages decide what screen is shown. Hooks fetch or manage data. Services call the backend. Lib files shape or format data. Components display data through props.",
    )

    heading(doc, "4. How The App Boots")
    for index, step in enumerate([
        "index.html provides <div id=\"root\"></div>.",
        "src/main.jsx creates the React root.",
        "src/App.jsx wraps everything in ApartmentsProvider.",
        "src/routes/AppRouter.jsx chooses the correct page for the URL.",
        "Pages call hooks. Hooks call services. Services call apiClient.js.",
        "The API response comes back, adapters normalize it, and components receive it as props.",
    ], start=1):
        number(doc, index, step)

    heading(doc, "5. Routing And Protection")
    add_table(
        doc,
        ["Route", "Component", "Access"],
        [
            ["/ and /discover", "src/pages/Discover.jsx", "Anyone"],
            ["/apartments", "src/pages/Neighborhoods.jsx", "Anyone"],
            ["/apartments/:id", "src/pages/ApartmentDetails.jsx", "Anyone"],
            ["/booking/:id", "src/pages/Booking.jsx", "Logged-in users"],
            ["/payment/:id", "src/pages/Payment.jsx", "Logged-in users"],
            ["/dashboard", "src/pages/Dashboard.jsx", "Logged-in users"],
            ["/history", "src/pages/Reservations.jsx", "Logged-in users"],
            ["/reservations/:id", "src/pages/ReservationDetails.jsx", "Logged-in users"],
            ["/profile", "src/pages/Profile.jsx", "Logged-in users"],
            ["/login", "src/pages/Login.jsx", "Logged-out users only"],
            ["/register", "src/pages/Register.jsx", "Logged-out users only"],
            ["/admin/*", "src/components/adminpage/*", "Admin users only"],
        ],
    )
    para(doc, "ProtectedRoute checks localStorage for access_token. PublicOnlyRoute keeps logged-in users away from Login/Register. Admin routes pass allowedRoles={[\"admin\"]}.")

    heading(doc, "6. Authentication And Tokens")
    para(doc, "A JWT token is a signed string from the backend. The app stores it in localStorage under access_token. apiClient.js automatically sends it as Authorization: Bearer <token> on later requests.")
    add_table(
        doc,
        ["File", "Responsibility"],
        [
            ["src/hooks/useAuthToken.js", "getStoredAuthToken, getStoredUserRole, saveToken, clearToken."],
            ["src/hooks/useUserRole.js", "Returns { role, isAdmin } based on the JWT payload."],
            ["src/services/authApi.js", "loginUser, registerUser, enableTwoFactor, confirmTwoFactor."],
            ["src/pages/Login.jsx", "Builds login body, saves token, redirects admin to /admin and guest to /discover."],
            ["src/components/Navbar.jsx", "Logout clears token and sends the user back to /discover or login flow."],
        ],
    )
    add_callout(
        doc,
        "Why Invalid or expired token appears",
        "The token may be old, copied incorrectly, from a normal user instead of an admin, or the backend may have invalidated it. Log in again, copy the new token, and confirm Postman uses Bearer Token without adding the word Bearer inside the token field.",
    )

    heading(doc, "7. API Layer")
    para(doc, "All backend calls should go through src/services/apiClient.js. That file owns the base URL, headers, token, JSON body, and error handling.")
    code_block(
        doc,
        """export async function apiRequest(path, options = {}) {
  const token = localStorage.getItem("access_token");
  const response = await fetch(`${API_BASE_URL}${path}`, {
    method: options.method || "GET",
    credentials: "include",
    headers: {
      Accept: "application/json",
      ...(options.body ? { "Content-Type": "application/json" } : {}),
      ...(token ? { Authorization: `Bearer ${token}` } : {}),
      ...options.headers,
    },
    body: options.body ? JSON.stringify(options.body) : undefined,
  });
  const payload = await response.json();
  if (!response.ok || payload.success === false) throw new Error(...);
  return payload.data;
}""",
    )
    add_table(
        doc,
        ["Service file", "Main endpoints"],
        [
            ["authApi.js", "/auth/login, /auth/register, /auth/2fa/enable, /auth/2fa/confirm"],
            ["propertiesApi.js", "GET /properties, GET /properties/:id"],
            ["bookingsApi.js", "GET /users/me/bookings, GET /admin/bookings, POST /bookings, POST /bookings/:id/cancel"],
            ["paymentsApi.js", "POST /payments/initiate"],
            ["usersApi.js", "GET/PATCH /users/me, PATCH /users/me/password, notifications"],
            ["adminApi.js", "Admin users, properties, bookings, tickets, permissions, config, audit logs"],
        ],
    )

    heading(doc, "8. Hooks And Context")
    add_table(
        doc,
        ["Hook or provider", "What it does"],
        [
            ["useProperties(params)", "Loads public properties and returns { properties, isLoading, error }."],
            ["useApartmentResource(id)", "Finds a static apartment first, otherwise loads GET /properties/:id."],
            ["useBookings(params)", "Loads booking history and exposes cancelBookingById and confirmAndCancelBooking."],
            ["useProfile()", "Loads user profile, notifications, 2FA, save profile, change password."],
            ["useAdminUsers()", "Loads users and exposes search, role update, active/inactive toggle."],
            ["useAdminPortal()", "Reads AdminPortalProvider state for sidebar, permissions, role counts, audit export."],
            ["useAdminTickets()", "Reads AdminTicketProvider state for support ticket list, replies, and status."],
            ["ApartmentsProvider", "Keeps the static fallback apartment catalog available."],
            ["AdminPortalProvider", "Wraps /admin so admin components share state."],
        ],
    )

    heading(doc, "9. Adapters And Data Shape")
    para(doc, "Adapters keep the UI stable even when the backend shape changes. For example, a backend property has id, title, neighborhood, price_per_night, max_guests, bedrooms, bathrooms, amenities, and images. propertyAdapter.js converts that into the apartment shape used by ApartmentCard, Booking, Payment, and detail pages.")
    add_table(
        doc,
        ["File", "Why it matters"],
        [
            ["src/lib/propertyAdapter.js", "Normalizes backend properties into UI apartments."],
            ["src/lib/bookingAdapter.js", "Normalizes backend bookings and computes dashboard stats/cancel rules."],
            ["src/lib/images.js", "Chooses backend image or fallback property image."],
            ["src/lib/apartmentFilters.js", "Filters by search, neighborhood, guests, category, and amenity."],
            ["src/lib/format.js", "Currency, dates, nights, capitalization."],
            ["src/lib/adminRoles.js", "Admin role and permission fallback definitions."],
        ],
    )

    doc.add_page_break()
    heading(doc, "10. Public And Guest Pages")
    add_table(
        doc,
        ["Page", "Main job", "Where to change it"],
        [
            ["Discover.jsx", "Landing page, hero, featured listings, neighborhoods.", "Discover.jsx and components/discover/*"],
            ["Neighborhoods.jsx", "Backend properties list and filters.", "useProperties.js, apartmentFilters.js, ApartmentCard.jsx"],
            ["ApartmentDetails.jsx", "Single property detail page.", "useApartmentResource.js and components/apartment-details/*"],
            ["Booking.jsx", "Create booking request.", "Booking.jsx and bookingsApi.js"],
            ["Payment.jsx", "Initiate payment.", "Payment.jsx and paymentsApi.js"],
            ["Dashboard.jsx", "User booking stats and upcoming reservations.", "Dashboard.jsx and components/dashboard/*"],
            ["Reservations.jsx", "Booking history and cancel booking.", "Reservations.jsx, useBookings.js, bookingAdapter.js"],
            ["ReservationDetails.jsx", "One booking detail view.", "ReservationDetails.jsx and bookingAdapter.js"],
            ["Profile.jsx", "Profile, password, 2FA, notifications.", "useProfile.js and components/profile/*"],
            ["Saved.jsx", "Saved places page still static.", "Replace local savedPlaces when backend exists"],
        ],
    )

    heading(doc, "11. Admin Portal")
    para(doc, "The admin portal lives under src/components/adminpage/ and is mounted by nested /admin routes in AppRouter.jsx. This keeps the admin system visually and structurally separate from the guest booking site.")
    add_table(
        doc,
        ["Admin route", "Component", "Backend status"],
        [
            ["/admin", "AdminDashboard.jsx", "Connected to admin totals."],
            ["/admin/users", "AdminUsers.jsx", "Connected to list, search/filter, role update, status update."],
            ["/admin/properties", "AdminProperties.jsx", "Connected to list properties by status."],
            ["/admin/bookings", "AdminBooking.jsx", "Connected to admin booking list."],
            ["/admin/content", "AdminContent.jsx", "Partly connected to config/content-style data; some controls are placeholders."],
            ["/admin/security", "AdminSettingss.jsx", "Connected to permissions through AdminPortalProvider."],
            ["/admin/support", "AdminSupport.jsx", "Connected to ticket list, detail, reply, and status."],
        ],
    )
    add_callout(
        doc,
        "Admin page build pattern",
        "Add the API function in adminApi.js, create a hook if loading/error state is needed, then pass clean props into the admin component. Keep admin-only shared state in AdminPortalProvider only when several admin pages need it.",
    )

    doc.add_page_break()
    heading(doc, "12. Reusable Components And Props")
    para(doc, "Reusable components receive data through props. This is how the app avoids repeating card layouts, status messages, password fields, apartment summaries, and dashboard widgets.")
    add_table(
        doc,
        ["Component", "Used for"],
        [
            ["Navbar.jsx", "Top navigation, user menu, logout."],
            ["MainLayout.jsx", "Wraps public pages with Navbar and Outlet."],
            ["PortalLayout.jsx", "Wraps dashboard/history/saved pages with the side menu."],
            ["ApartmentCard.jsx", "Property card on /apartments and featured sections."],
            ["ApartmentSummary.jsx", "Booking and payment summary panel."],
            ["FilterChipGroup.jsx", "Neighborhood and amenity chip filters."],
            ["StatusMessage.jsx", "Reusable loading/error/success feedback."],
            ["PasswordField.jsx", "Show/hide password input for Login/Register/Profile."],
            ["LoadingScreen.jsx", "Full-page loading state."],
        ],
    )

    heading(doc, "13. Styling With Tailwind")
    para(doc, "Tailwind classes are written directly inside JSX. The app uses #F7F8F0 for the soft page background, sky-900/sky-950 for primary brand actions, slate colors for readable text, and light borders for cards.")
    add_table(
        doc,
        ["File", "What to edit"],
        [
            ["tailwind.config.js", "Theme extension, content globs, custom colors if needed."],
            ["src/index.css", "Tailwind directives only."],
            ["Individual JSX files", "Most spacing, color, layout, borders, and responsive classes."],
        ],
    )
    add_callout(
        doc,
        "Changing colors",
        "Start by changing repeated Tailwind utilities such as bg-sky-900, text-sky-900, border-[#E7E8DE], and bg-[#F7F8F0]. Use global search so you see every place the brand color appears.",
    )

    heading(doc, "14. Backend Coverage")
    add_table(
        doc,
        ["Area", "Status"],
        [
            ["Authentication", "Connected"],
            ["Public properties", "Connected"],
            ["Property details", "Connected, with static fallback"],
            ["Booking creation", "Connected"],
            ["Cancel booking", "Connected"],
            ["Payment initiation", "Connected"],
            ["Dashboard/history/reservation details", "Connected to booking endpoints"],
            ["Profile/password/2FA", "Connected"],
            ["Admin users/properties/bookings/support/security", "Connected or partly connected"],
            ["Saved places", "Static"],
            ["Public support page", "Static"],
            ["Some admin content editor actions", "Partly placeholder"],
        ],
    )

    heading(doc, "15. Common Change Recipes")
    recipes = [
        ("Add a new route", [
            "Create a page or component file.",
            "Import it in src/routes/AppRouter.jsx.",
            "Add a <Route> with the correct path.",
            "Wrap it in ProtectedRoute when login is required.",
            "Add a Navbar or AdminLayout link if users need to reach it.",
        ]),
        ("Add a backend endpoint", [
            "Create or update a service file in src/services/.",
            "Call apiRequest or apiGetList.",
            "If the data shape is new, normalize it in src/lib/.",
            "Use it from a hook when the UI needs loading/error state.",
        ]),
        ("Add a property in the backend", [
            "Log in as admin and get a fresh admin JWT.",
            "POST /properties with the property body.",
            "If status is draft, PATCH /properties/:id with { \"status\": \"published\" }.",
            "Refresh /apartments and confirm GET /properties returns it.",
        ]),
        ("Archive a property", [
            "Use PATCH /properties/:id with { \"status\": \"archived\" } if supported.",
            "Public /apartments should only show published listings.",
            "Admin /admin/properties can still show draft or archived listings by status filter.",
        ]),
        ("Connect Saved places later", [
            "Add saved/favorites functions to a new or existing service.",
            "Create useSavedPlaces() for loading/error/cancel or remove actions.",
            "Replace savedPlaces inside Saved.jsx with the hook output.",
        ]),
        ("Add a new admin page", [
            "Create src/components/adminpage/AdminSomething.jsx.",
            "Add a nested route under /admin in AppRouter.jsx.",
            "Add a sidebar link in the nav array inside AdminLayout.jsx.",
            "Connect backend calls through adminApi.js.",
        ]),
    ]
    for title, steps in recipes:
        heading(doc, title, 3)
        for index, step in enumerate(steps, start=1):
            number(doc, index, step)

    heading(doc, "16. Debug Checklist")
    add_table(
        doc,
        ["Problem", "Likely cause", "What to check"],
        [
            ["Invalid or expired token", "Old/missing token or wrong role.", "Log in again, inspect localStorage access_token, confirm role payload."],
            ["Property created but not visible", "It is draft or not returned by public endpoint.", "GET /properties and confirm status is published."],
            ["CORS error", "Backend blocked browser origin.", "Backend CORS settings, not React code."],
            ["401 in Postman", "Bearer token not sent correctly.", "Authorization tab: Bearer Token, token only."],
            ["Page shows fallback image", "Backend images array is empty.", "Add image URLs or update FALLBACK_PROPERTY_IMAGE."],
            ["Static data still visible", "Screen not connected yet or fallback is active.", "Use this guide section 14 to identify status."],
        ],
    )

    heading(doc, "17. Glossary")
    add_table(
        doc,
        ["Term", "Plain meaning"],
        [
            ["Component", "A reusable UI function that returns JSX."],
            ["Props", "Inputs passed from parent to child component."],
            ["State", "Data a component owns and can change."],
            ["Hook", "A reusable function starting with use that manages state, effects, or context."],
            ["Context", "Shared app state available without passing props through every layer."],
            ["Service", "A file that wraps one group of backend endpoints."],
            ["Adapter", "A helper that converts backend data to UI-friendly data."],
            ["JWT", "The login token saved in localStorage and sent with API requests."],
            ["ProtectedRoute", "A router wrapper that blocks users who are not logged in or not allowed."],
            ["Tailwind", "CSS utility classes written directly in JSX."],
        ],
    )

    doc.add_section(WD_SECTION.CONTINUOUS)
    footer = doc.sections[-1].footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run("Urban Sanctuary Booking Portal - Beginner Developer Guide")
    run.font.size = Pt(8)
    run.font.color.rgb = MUTED

    doc.save(OUTPUT)


if __name__ == "__main__":
    build_doc()
    print(f"Wrote {OUTPUT}")
